"""
Service d'extraction de texte depuis différents formats de documents
"""
import logging
from io import BytesIO
from typing import Optional
import PyPDF2
from docx import Document
from config import settings

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extrait le texte d'un fichier PDF
    
    Args:
        file_content: Contenu binaire du fichier PDF
    
    Returns:
        Texte extrait
    """
    try:
        pdf_file = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        extracted_text = "\n".join(text_parts)
        
        # Si le PDF est scanné (peu de texte), utiliser OCR
        if len(extracted_text.strip()) < 100 and settings.OCR_ENABLED:
            logger.info("[OCR] PDF semble scanne, activation de l'OCR...")
            extracted_text = extract_text_with_ocr(file_content)
        
        return extracted_text
        
    except Exception as e:
        logger.error(f"Erreur extraction PDF: {str(e)}")
        # Fallback vers OCR si disponible
        if settings.OCR_ENABLED:
            logger.info("[OCR] Tentative extraction via OCR...")
            return extract_text_with_ocr(file_content)
        raise


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extrait le texte d'un fichier DOCX
    
    Args:
        file_content: Contenu binaire du fichier DOCX
    
    Returns:
        Texte extrait
    """
    try:
        docx_file = BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        
        # Extraire les paragraphes
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extraire les tableaux
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"Erreur extraction DOCX: {str(e)}")
        raise


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extrait le texte d'un fichier TXT
    
    Args:
        file_content: Contenu binaire du fichier TXT
    
    Returns:
        Texte extrait
    """
    try:
        # Tenter différents encodages
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                return text
            except UnicodeDecodeError:
                continue
        
        # Si aucun encodage ne fonctionne
        raise ValueError("Impossible de décoder le fichier TXT")
        
    except Exception as e:
        logger.error(f"Erreur extraction TXT: {str(e)}")
        raise


def extract_text_with_ocr(file_content: bytes) -> str:
    """
    Extrait le texte via OCR (Tesseract)
    
    Args:
        file_content: Contenu binaire du fichier image ou PDF scanné
    
    Returns:
        Texte extrait via OCR
    """
    try:
        import pytesseract
        from PIL import Image
        import pdf2image
        
        # Configurer le chemin Tesseract si spécifié
        if settings.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
        
        # Convertir PDF en images si nécessaire
        try:
            images = pdf2image.convert_from_bytes(file_content)
        except:
            # Si ce n'est pas un PDF, essayer de l'ouvrir comme image
            images = [Image.open(BytesIO(file_content))]
        
        text_parts = []
        for i, image in enumerate(images):
            logger.info(f"[OCR] OCR page {i+1}/{len(images)}...")
            text = pytesseract.image_to_string(image, lang='fra')
            if text.strip():
                text_parts.append(text)
        
        return "\n".join(text_parts)
        
    except ImportError:
        logger.warning("[WARN] OCR non disponible (pytesseract ou pdf2image manquant)")
        return ""
    except Exception as e:
        logger.error(f"Erreur OCR: {str(e)}")
        return ""


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extrait le texte d'un fichier selon son extension
    
    Args:
        file_content: Contenu binaire du fichier
        filename: Nom du fichier (pour déterminer l'extension)
    
    Returns:
        Texte extrait
    """
    file_extension = f".{filename.split('.')[-1].lower()}"
    
    logger.info(f"[EXTRACTION] Extraction du texte de {filename} (type: {file_extension})")
    
    if file_extension == ".pdf":
        return extract_text_from_pdf(file_content)
    elif file_extension in [".docx", ".doc"]:
        return extract_text_from_docx(file_content)
    elif file_extension == ".txt":
        return extract_text_from_txt(file_content)
    else:
        raise ValueError(f"Type de fichier non supporté: {file_extension}")


def extract_text_with_tika(file_content: bytes) -> Optional[str]:
    """
    Extrait le texte via Apache Tika (méthode alternative)
    
    Args:
        file_content: Contenu binaire du fichier
    
    Returns:
        Texte extrait ou None si erreur
    """
    try:
        from tika import parser
        
        # Parser le document
        if settings.TIKA_SERVER_URL:
            parsed = parser.from_buffer(file_content, serverEndpoint=settings.TIKA_SERVER_URL)
        else:
            parsed = parser.from_buffer(file_content)
        
        return parsed.get('content', '')
        
    except ImportError:
        logger.warning("[WARN] Tika non disponible")
        return None
    except Exception as e:
        logger.error(f"Erreur Tika: {str(e)}")
        return None
